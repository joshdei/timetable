#!/usr/bin/env python3
"""
Timetable Extractor: Counts tables, shows locations, and extracts/displays data stored in tables from DOCX.
Usage: python timetable_extractor.py [--input file.docx]
Requires: python-docx library (pip install python-docx)
"""

from docx import Document
import sys
import os
import argparse

def print_table_data(table, table_index):
    """Extract and print data from a table."""
    print(f"\n--- Table {table_index} Data ---")
    rows, cols = len(table.rows), 0
    if table.rows:
        cols = len(table.rows[0].cells)
    
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for cell_idx, cell in enumerate(row.cells):
            cell_text = ' '.join([paragraph.text for paragraph in cell.paragraphs]).strip()
            row_data.append(cell_text)
        print(f"Row {row_idx}: {' | '.join(row_data)}")
    print(f"Table {table_index} dimensions: {rows} rows x {cols} columns")

def main():
    parser = argparse.ArgumentParser(description="Extract tables from DOCX timetable")
    parser.add_argument("--input", "-i", default=r"c:/Users/Devaris/Desktop/timetbale/timetable.docx",
                        help="Input DOCX file (default: timetable.docx)")
    args = parser.parse_args()
    
    docx_path = args.input
    
    if not os.path.exists(docx_path):
        print(f"Error: File not found at {docx_path}")
        sys.exit(1)
    
    try:
        doc = Document(docx_path)
        tables = doc.tables
        table_count = len(tables)
        
        print(f"Document: {docx_path}")
        print(f"Total number of tables: {table_count}")
        
        if table_count == 0:
            print("No tables found in the document.")
            print("Document text content:")
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            print(full_text[:2000] + ("..." if len(full_text) > 2000 else ""))
        else:
            print("\nTable locations (indices) and data:")
            for i, table in enumerate(tables):
                print_table_data(table, i)
        
        print("\nExtraction complete.")
        
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading DOCX: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()

