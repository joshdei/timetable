#!/usr/bin/env python3
"""
Remove empty rows from timetable.docx table. Saves cleaned version.
"""
from docx import Document
import os
import re

def is_row_empty(row):
    """
    Check if all cells in row are empty/whitespace. Returns True if empty.
    """
    for cell in row.cells:
        cell_text = ' '.join(p.text.strip() for p in cell.paragraphs).strip()
        if cell_text and not re.match(r'^\\s*$', cell_text):
            return False
    return True

def main():
    input_path = r'c:/Users/Devaris/Desktop/timetbale/timetable.docx'
    output_path = r'c:/Users/Devaris/Desktop/timetbale/cleaned_timetable.docx'
    
    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found.")
        return 1
    
    doc = Document(input_path)
    tables = doc.tables
    
    if not tables:
        print("No tables found.")
        return 1
    
    # Target main table (index 0)
    table = tables[0]
    original_rows = len(table.rows)
    
    # Delete from end to preserve indices
    rows_to_delete = 0
    for row_idx in range(len(table.rows) - 1, -1, -1):
        if is_row_empty(table.rows[row_idx]):
            del table.rows[row_idx]
            rows_to_delete += 1
    
    doc.save(output_path)
    print(f"Cleaned: Removed {rows_to_delete} empty rows.")
    print(f"Original: {original_rows} rows -> Final: {len(table.rows)} rows")
    print(f"Saved to: {output_path}")
    return 0

if __name__ == "__main__":
    import sys
    sys.exit(main())

