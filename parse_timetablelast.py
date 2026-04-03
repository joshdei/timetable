#!/usr/bin/env python3
"""
Align timetable subjects with time slots and dates. Outputs structured JSON/CSV.
"""
from docx import Document
import json
import csv
import re
import os

# Time slot labels to filter out junk header rows
TIME_LABELS = {'8:30am', '11am', '2pm', '8:30 am', '11 am', '2 pm'}

def is_header_row(entry):
    """Return True if this row is a label/header row (code equals a time string)."""
    for slot in ['8_30am', '11am', '2pm']:
        if entry[slot] and entry[slot]['code'].strip().lower() in TIME_LABELS:
            return True
    return False

def parse_timetable(doc_path):
    doc = Document(doc_path)
    table = doc.tables[0]  # Main table
    rows = []

    current_day = None
    current_day_name = None  # FIX 1: track day name across rows, just like current_day
    current_date = None

    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = ' '.join(p.text.strip() for p in cell.paragraphs).strip()
            cells.append(cell_text)

        if len(cells) < 8:
            continue

        # Update current day number when we see a digit in the first cell
        day_cell = cells[0].strip()
        if day_cell.isdigit() and len(day_cell) == 1:
            current_day = int(day_cell)

        # FIX 2: corrected regex — no double-backslash inside raw string
        date_match = re.search(r'(\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]+,?\s+\d{4})', ' '.join(cells))
        if date_match:
            current_date = date_match.group(1).strip()

        # FIX 1 (cont): update day_name only when cells[1] is a meaningful non-date string
        candidate_name = cells[1].strip() if len(cells) > 1 else ''
        if candidate_name and not date_match:
            # Only update if it looks like a weekday name, not a date string
            if not re.search(r'\d', candidate_name):
                current_day_name = candidate_name

        entry = {
            'day': current_day,
            'day_name': current_day_name,   # FIX 1: use persisted value, not raw cells[1]
            'date': current_date,
            '8_30am': {'code': cells[2], 'title': cells[3]} if cells[2] or cells[3] else None,
            '11am':   {'code': cells[4], 'title': cells[5]} if cells[4] or cells[5] else None,
            '2pm':    {'code': cells[6], 'title': cells[7]} if cells[6] or cells[7] else None,
        }

        # Only keep rows that have at least one subject slot populated
        if not any(entry[slot] for slot in ['8_30am', '11am', '2pm']):
            continue

        # FIX 3: drop junk header rows where the code IS the time label
        if is_header_row(entry):
            continue

        rows.append(entry)

    # BACKFILL: first row(s) of a day may have date=None because the date
    # appears later in the same day group. Find the date for each day, then
    # apply it backwards to any rows that missed it.
    day_dates = {}
    for entry in rows:
        if entry['date'] and entry['day'] not in day_dates:
            day_dates[entry['day']] = entry['date']

    for entry in rows:
        if entry['date'] is None and entry['day'] in day_dates:
            entry['date'] = day_dates[entry['day']]

    return rows

TIME_SLOT_LABELS = {
    '8_30am': '8:30am',
    '11am':   '11am',
    '2pm':    '2pm',
}

def flatten_to_model(rows, exam_type):
    """
    Flatten parsed rows into one record per subject, matching the
    Laravel Timetable model columns:
      exam_date | type_of_time_table | time_slot | course_code | course_title | status
    """
    records = []
    for row in rows:
        for slot_key, slot_label in TIME_SLOT_LABELS.items():
            subject = row.get(slot_key)
            if not subject:
                continue
            records.append({
                'exam_date':          row['date'],
                'type_of_time_table': exam_type,
                'time_slot':          slot_label,
                'course_code':        subject['code'],
                'course_title':       subject['title'],
                'status':             'active',
            })
    return records

def save_json(data, path):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def save_csv(data, path):
    fieldnames = [
        'exam_date', 'type_of_time_table', 'time_slot',
        'course_code', 'course_title', 'status',
    ]
    with open(path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in data:
            writer.writerow(row)

def main():
    input_path = r'c:/Users/Devaris/Desktop/timetbale/timetable/cleaned_timetable.docx'
    json_path  = r'c:/Users/Devaris/Desktop/timetbale/aligned_timetable.json'
    csv_path   = r'c:/Users/Devaris/Desktop/timetbale/aligned_timetable.csv'

    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found. Run remove_empty_rows.py first.")
        return

    # Prompt user for exam type
    print("📝 What type of exam timetable is this?")
    print("   e.g. Final Exam, Mid-Semester, Supplementary, Quiz, etc.")
    exam_type = input("   Enter exam type: ").strip()
    if not exam_type:
        print("❌ Exam type cannot be empty. Exiting.")
        return

    raw_data = parse_timetable(input_path)
    data     = flatten_to_model(raw_data, exam_type)

    save_json(data, json_path)
    save_csv(data, csv_path)

    print(f"\n✅ Exported {len(data)} subject records!")
    print(f"📄 JSON: {json_path}")
    print(f"📊 CSV:  {csv_path}")

    print("\n📋 Sample records:")
    for record in data[:3]:
        print(f"  [{record['time_slot']}] {record['course_code']} — {record['course_title']}")
        print(f"         Date: {record['exam_date']}  |  Type: {record['type_of_time_table']}  |  Status: {record['status']}")
    print()

if __name__ == "__main__":
    main()
