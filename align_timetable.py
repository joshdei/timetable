#!/usr/bin/env python3
"""
Align timetable subjects with time slots and dates. Outputs structured JSON/CSV.
"""
from docx import Document
import json
import csv
import re
import os

def parse_timetable(doc_path):
    doc = Document(doc_path)
    table = doc.tables[0]  # Main table
    rows = []
    
    current_day = None
    current_date = None
    
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = ' '.join(p.text.strip() for p in cell.paragraphs).strip()
            cells.append(cell_text)
        
        if len(cells) < 8:
            continue
            
        day_cell = cells[0].strip()
        if day_cell.isdigit() and len(day_cell) == 1:
            current_day = int(day_cell)
        
        date_match = re.search(r'(\\d{1,2}(?:st|nd|rd|th)? [A-Za-z, ]+ \\d{4})', ' '.join(cells))
        if date_match:
            current_date = date_match.group(1)
        
        day_name = cells[1].strip() if len(cells) > 1 else ''
        
        entry = {
            'day': current_day,
            'day_name': day_name,
            'date': current_date,
            '8_30am': {'code': cells[2], 'title': cells[3]} if cells[2] or cells[3] else None,
            '11am': {'code': cells[4], 'title': cells[5]} if cells[4] or cells[5] else None,
            '2pm': {'code': cells[6], 'title': cells[7]} if cells[6] or cells[7] else None
        }
        if any(entry[slot] for slot in ['8_30am', '11am', '2pm']):
            rows.append(entry)
    
    return rows

def save_json(data, path):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def save_csv(data, path):
    fieldnames = ['day', 'day_name', 'date', '8_30am_code', '8_30am_title', '11am_code', '11am_title', '2pm_code', '2pm_title']
    with open(path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in data:
            writer.writerow({
                'day': row['day'],
                'day_name': row['day_name'],
                'date': row['date'],
                '8_30am_code': row['8_30am']['code'] if row['8_30am'] else '',
                '8_30am_title': row['8_30am']['title'] if row['8_30am'] else '',
                '11am_code': row['11am']['code'] if row['11am'] else '',
                '11am_title': row['11am']['title'] if row['11am'] else '',
                '2pm_code': row['2pm']['code'] if row['2pm'] else '',
                '2pm_title': row['2pm']['title'] if row['2pm'] else ''
            })

def main():
    input_path = r'c:/Users/Devaris/Desktop/timetbale/cleaned_timetable.docx'
    json_path = r'c:/Users/Devaris/Desktop/timetbale/aligned_timetable.json'
    csv_path = r'c:/Users/Devaris/Desktop/timetbale/aligned_timetable.csv'
    
    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found. Run remove_empty_rows.py first.")
        return
    
    data = parse_timetable(input_path)
    
    save_json(data, json_path)
    save_csv(data, csv_path)
    
    print(f"✅ Aligned {len(data)} timetable entries!")
    print(f"📄 JSON: {json_path}")
    print(f"📊 CSV: {csv_path}")
    
    print("\n📋 Sample entries:")
    for entry in data[:2]:
        print(f"Day {entry['day']} ({entry['day_name']}, {entry['date']}):")
        for slot in ['8_30am', '11am', '2pm']:
            if entry[slot]:
                print(f"  {slot}: {entry[slot]['code']} - {entry[slot]['title']}")
        print()
    
    total_subjects = sum(1 for entry in data for slot in ['8_30am', '11am', '2pm'] if entry[slot])
    print(f"🎯 Total subjects aligned: {total_subjects}")

if __name__ == "__main__":
    main()

