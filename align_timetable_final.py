#!/usr/bin/env python3
# FINAL - Simple, fast timetable alignment
from docx import Document
import json
import os
import re

input_path = r'c:/Users/Devaris/Desktop/timetbale/cleaned_timetable.docx'
doc = Document(input_path)
table = doc.tables[0]

current_header = ""
schedule = []

data = []

for row in table.rows:
    cells = [''.join([p.text.strip() for p in cell.paragraphs]).strip() for cell in row.cells]
    
    if len(cells) < 8:
        continue
    
    # Day header (digits in col0)
    if re.match(r'^\\d+$', cells[0]):
        if schedule:
            data.append({
                'header': current_header,
                'schedule': schedule
            })
            schedule = []
        current_header = cells[0] + ' ' + cells[1]
    else:
        # Subject row
        if cells[2]:
            schedule.append('8:30am - ' + cells[2] + ' - ' + cells[3])
        if cells[4]:
            schedule.append('11am - ' + cells[4] + ' - ' + cells[5])
        if cells[6]:
            schedule.append('2pm - ' + cells[6] + ' - ' + cells[7])

# Last entry
if schedule:
    data.append({
        'header': current_header,
        'schedule': schedule
    })

# Save
with open('timetable_final.json', 'w') as f:
    json.dump(data, f, indent=2)

print('✅ COMPLETE! timetable_final.json created')
print(f'{len(data)} days aligned')
print('\\nSample:')
print(json.dumps(data[0], indent=2))

